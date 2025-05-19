import os
import time
import docx
from docx import Document
from typing import Dict, Any, Optional
import tempfile
import config

class DocumentProcessor:
    """문서 생성 및 처리를 위한 유틸리티 클래스"""
    
    @staticmethod
    def create_docx(content: str, title: str) -> str:
        """마크다운 텍스트로부터 Word 문서 생성"""
        doc = Document()
        
        # 제목 추가
        doc.add_heading(title, level=1)
        
        # 내용 추가 (마크다운 파싱은 간단하게 처리)
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                if para.startswith('# '):
                    doc.add_heading(para[2:], level=1)
                elif para.startswith('## '):
                    doc.add_heading(para[3:], level=2)
                elif para.startswith('### '):
                    doc.add_heading(para[4:], level=3)
                else:
                    doc.add_paragraph(para)
        
        # 임시 파일 이름 생성
        timestamp = int(time.time())
        filename = f"{config.TEMP_DIR}/{title.replace(' ', '_')}_{timestamp}.docx"
        
        # 문서 저장
        doc.save(filename)
        
        return filename
    
    @staticmethod
    def get_document_title(document_type: str, user_input: Dict[str, Any]) -> str:
        """문서 유형과 사용자 입력에 따른 문서 제목 생성"""
        
        if document_type == "contract":
            contract_type = user_input.get('contract_type', '계약')
            party_a = user_input.get('party_a', '갑')
            party_b = user_input.get('party_b', '을')
            return f"{contract_type} 계약서 ({party_a}-{party_b})"
            
        elif document_type == "litigation":
            case_type = user_input.get('case_type', '소송')
            subtype = user_input.get('document_subtype', '소장')
            return f"{case_type} {subtype}"
            
        elif document_type == "legal_opinion":
            topic = user_input.get('legal_question', '')
            if len(topic) > 30:
                topic = topic[:30] + '...'
            return f"법률 의견서 - {topic}"
            
        elif document_type == "will":
            testator = user_input.get('testator', '유언자')
            will_type = user_input.get('will_type', '유언장')
            return f"{testator}의 {will_type}"
            
        elif document_type == "corporate":
            company = user_input.get('company_name', '회사')
            doc_type = user_input.get('document_subtype', '기업 문서')
            return f"{company} {doc_type}"
            
        elif document_type == "real_estate":
            prop_type = user_input.get('property_type', '부동산')
            doc_type = user_input.get('document_subtype', '계약서')
            return f"{prop_type} {doc_type}"
            
        elif document_type == "ip":
            ip_type = user_input.get('ip_type', '지식재산권')
            doc_type = user_input.get('document_subtype', '계약서')
            return f"{ip_type} {doc_type}"
            
        elif document_type == "compliance":
            area = user_input.get('regulatory_area', '규제')
            org = user_input.get('organization', '조직')
            return f"{org} {area} 준수 문서"
            
        else:
            return f"법률 문서 {int(time.time())}"
    
    @staticmethod
    def format_document_preview(content: str) -> str:
        """Streamlit 표시를 위한 문서 미리보기 포맷팅"""
        # 이미 마크다운 형식이므로 그대로 반환
        return content